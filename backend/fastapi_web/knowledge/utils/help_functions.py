"""Вспомогательные функции приложения Знания."""
import base64
import hashlib
import io
import json
import mimetypes
import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple, Union

import aiofiles
import docx
import httpx
import openpyxl
import pdfplumber
from bs4 import BeautifulSoup
from fastapi import HTTPException, UploadFile

from chats.db.mongo.schemas import ChatMessage
from db.mongo.db_init import mongo_db
from db.redis.db_init import redis_db
from gemini_base.gemini_init import gemini_client
from infra import settings
from knowledge.db.mongo.enums import ContextPurpose, ContextType
from knowledge.db.mongo.schemas import ContextEntry, KnowledgeBase
from knowledge.utils.prompts import AI_PROMPTS
from openai_base.openai_init import openai_client
from utils.help_functions import split_prompt_parts

from .knowledge_base import KNOWLEDGE_BASE

# ==============================
# КОНСТАНТЫ
# ==============================

IMG_EXT = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
DOC_EXT = {".pdf", ".docx", ".doc", ".xlsx", ".xls"}


# ==============================
# БЛОК: СРАВНЕНИЕ И СЛИЯНИЕ
# ==============================


def mark_created_diff(val: Any) -> Any:
    """Помечает новые данные тегом `created`."""
    if isinstance(val, dict):
        return {"tag": "created", **
                {k: mark_created_diff(v) for k, v in val.items()}}
    return val


def compare_values(old_value, new_value) -> Optional[dict]:
    if isinstance(old_value, dict) and isinstance(new_value, dict):
        nested = diff_with_tags(old_value, new_value)
        return {"tag": "updated", **nested} if nested else None
    return (
        {"tag": "updated", "old": old_value, "new": new_value}
        if old_value != new_value
        else None
    )


def diff_with_tags(old: dict, new: dict) -> Optional[dict]:
    """Словарь различий с CRUD-тегами."""
    diff = {}
    for k in old.keys() | new.keys():
        if k not in old:
            diff[k] = mark_created_diff(new[k])
        elif k not in new:
            diff[k] = {"tag": "deleted", "old": old[k]}
        else:
            nested = compare_values(old[k], new[k])
            if nested:
                diff[k] = nested
    return diff or None


def deep_merge(original: dict, patch: dict) -> dict:
    """Рекурсивное слияние с поддержкой `_delete`."""
    if not (isinstance(original, dict) and isinstance(patch, dict)):
        return patch
    result = original.copy()
    for k, v in patch.items():
        if isinstance(v, dict) and v.get("_delete"):
            result.pop(k, None)
        elif isinstance(v, dict):
            result[k] = deep_merge(result.get(k, {}), v)
        else:
            result[k] = v
    return result


def merge_external_structures(main_kb: dict, externals: list[dict]) -> dict:
    """Накатывает внешние структуры поверх основной (приоритет справа)."""
    merged = deepcopy(main_kb)
    for ext in externals:
        merged = deep_merge(merged, ext)
    return merged


# ==============================
# БЛОК: БАЗА ЗНАНИЙ
# ==============================


async def get_knowledge_full_document() -> dict:
    """Документ БЗ «main» без `_id`."""
    doc = await mongo_db.knowledge_collection.find_one({"app_name": settings.APP_NAME})
    if not doc:
        raise HTTPException(404, "Knowledge base not found")
    doc.pop("_id", None)
    return doc


async def get_knowledge_base() -> Tuple[dict, Optional[KnowledgeBase]]:
    """(dict, pydantic-model|None)."""
    doc = await get_knowledge_full_document()
    return (doc, KnowledgeBase(
        **doc)) if doc.get("knowledge_base") is not None else (KNOWLEDGE_BASE, None)


async def save_kb_doc(data: dict) -> bool:
    """Сохраняет документ БЗ."""
    res = await mongo_db.knowledge_collection.replace_one(
        {"app_name": settings.APP_NAME},
        data,
        upsert=True,
    )
    return bool(res.modified_count or res.upserted_id)


# ==============================
# БЛОК: ПАРСИНГ
# ==============================


async def parse_pdf(file: UploadFile) -> str:
    """Возвращает текст из PDF-файла."""
    data = await file.read()
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        return "\n".join(page.extract_text()
                         or "" for page in pdf.pages).strip()


async def parse_docx(file: UploadFile) -> str:
    """Возвращает текст (параграфы + таблицы) из DOCX."""
    data = await file.read()
    doc = docx.Document(io.BytesIO(data))

    parts: list[str] = []

    # 1. Параграфы
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # 2. Таблицы (по строкам)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts).strip()


async def parse_excel(file: UploadFile) -> str:
    """Возвращает текст из Excel-файла (все листы, ячейки через пробел)."""
    data = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(data))
    return "\n".join(
        " ".join(str(c) if c else "" for c in row)
        for sh in wb.worksheets
        for row in sh.iter_rows(values_only=True)
    ).strip()


async def parse_file(upload_file: UploadFile) -> Optional[str]:
    """Определяет парсер по расширению и извлекает текст."""
    ext = Path(upload_file.filename).suffix.lower()
    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        # ".doc":  parse_doc,
        ".xlsx": parse_excel,
        ".xls": parse_excel,
    }
    return await parsers[ext](upload_file) if ext in parsers else None


async def parse_file_from_path(path: str) -> str:
    """Читает файл по пути и извлекает текст (DOC/PDF/Excel)."""
    ext = Path(path).suffix.lower()
    if ext not in DOC_EXT:
        return ""
    async with aiofiles.open(path, "rb") as f:
        fake = UploadFile(filename=Path(path).name, file=io.BytesIO(await f.read()))
    return await parse_file(fake) or ""


async def parse_url(url: str, timeout: int = 10) -> str:
    """Скачивает страницу и возвращает очищенный текст без скриптов/стилей."""
    url = url.strip()
    if not url:
        raise ValueError("Empty URL passed to parse_url")

    async with httpx.AsyncClient(
        timeout=timeout,
        follow_redirects=True,
        headers={"User-Agent": "Mozilla/5.0 (knowledge-context-bot)"},
    ) as client:
        r = await client.get(url)
        r.raise_for_status()

    soup = BeautifulSoup(r.text, "lxml")
    for t in soup(["script", "style", "noscript", "header", "footer", "form"]):
        t.decompose()
    return "\n".join(line.strip()
                     for line in soup.get_text("\n").splitlines() if line.strip())


async def cache_url_snapshot(url: str, ttl: int = settings.CONTEXT_TTL) -> str:
    """Возвращает кешированный текст страницы или кэширует новый."""
    url = url.strip()
    if not url:
        raise ValueError(
            "Empty or whitespace-only URL passed to cache_url_snapshot")

    key = f"url_cache:{hashlib.sha1(url.encode()).hexdigest()}"
    cached = await redis_db.get(key)
    if cached:
        return cached.decode(errors="ignore") if isinstance(
            cached, bytes) else cached

    parsed = await parse_url(url)
    await redis_db.set(key, parsed, ex=ttl)
    return parsed


# ---------- Сохранение загруженных файлов ----------

async def save_uploaded_file(file: UploadFile, uid: str) -> Path:
    """Сохраняет UploadFile на диск и возвращает путь."""
    dst_dir = Path(settings.CONTEXT_PATH) / uid
    dst_dir.mkdir(parents=True, exist_ok=True)
    dst_path = dst_dir / file.filename
    async with aiofiles.open(dst_path, "wb") as f:
        await f.write(await file.read())
    return dst_path

# ==============================
# БЛОК: LLM-ПОМOЩНИКИ
# ==============================


def pick_model_and_client(model: str):
    """Возвращает подходящий клиент и имя модели."""
    if model.startswith("gpt"):
        return openai_client, model
    if model.startswith("gemini"):
        return gemini_client, model
    return gemini_client, settings.MODEL_DEFAULT


async def admin_chat_generate_any(
    client,
    model: str,
    messages: list[dict],
    system_instruction: Optional[str] = None,
) -> str:
    """Отправляет сообщения в LLM и возвращает ответ."""
    if model.startswith("gpt"):
        resp = await client.chat.completions.create(model=model, messages=messages, temperature=0.1)
        return resp.choices[0].message.content.strip()

    resp = await client.chat_generate(
        model=model,
        messages=messages,
        temperature=0.1,
        system_instruction=system_instruction,
    )
    return resp["candidates"][0]["content"]["parts"][0]["text"].strip()


def extract_json_from_gpt(raw: str) -> dict:
    """Извлекает первый JSON из сырой строки ответа GPT."""
    match = re.search(r"\{.*\}", raw, re.DOTALL)
    if not match:
        return {}
    try:
        return json.loads(match.group(0))
    except json.JSONDecodeError:
        return {}


# ==============================
# БЛОК: СООБЩЕНИЯ, ФАЙЛЫ
# ==============================


def extract_english_value(value: Any) -> Any:
    """Возвращает 'en' значение из i18n-строки / словаря."""
    if isinstance(value, str):
        try:
            return json.loads(value).get("en", value)
        except (json.JSONDecodeError, TypeError):
            return value
    if isinstance(value, dict):
        return value.get("en", value)
    if isinstance(value, list):
        return [extract_english_value(v) for v in value]
    return value


def build_messages_for_model(
    system_prompt: Optional[str],
    messages_data: Union[List["ChatMessage"], List[dict]],
    user_message: str,
    model: str,
) -> Dict[str, Any]:
    """Формирует payload сообщений для конкретного LLM с поддержкой разделения статики и динамики."""
    messages, system_instruction = [], None

    static_part, dynamic_part = split_prompt_parts(system_prompt or "")

    if model.startswith("gpt"):
        if static_part:
            messages.append({"role": "system", "content": static_part})
        if dynamic_part:
            messages.append({"role": "system", "content": dynamic_part})
    elif model.startswith("gemini"):
        full_prompt = f"{static_part}\n\n{dynamic_part}".strip()
        messages.append({"role": "user", "parts": [{"text": full_prompt}]})
    else:
        full_prompt = f"{static_part}\n\n{dynamic_part}".strip()
        system_instruction = full_prompt

    for raw in messages_data:
        role_raw, content = ("user", raw)
        if "ChatMessage" in str(type(raw)):
            role_raw = extract_english_value(raw.sender_role.value)
            content = raw.message
        elif isinstance(raw, dict):
            role_raw = extract_english_value(raw.get("sender_role", "user"))
            content = raw

        role = "assistant" if role_raw.lower() in {
            "assistant", "consultant", "ai assistant"} else "user"
        blocks: List[Dict[str, Any]] = []

        if isinstance(content, str):
            blocks.append({"type": "text", "text": content})
        elif isinstance(content, dict):
            if "text" in content:
                blocks.append({"type": "text", "text": content["text"]})
            elif "image_url" in content:
                blocks.append(
                    {"type": "image_url", "image_url": content["image_url"]})

        if not blocks:
            continue

        if model.startswith("gpt"):
            messages.append({"role": role, "content": blocks})
        else:
            parts = [
                {"text": b["text"]}
                if b["type"] == "text"
                else {
                    "inline_data": {
                        "mime_type": "image/jpeg",
                        "data": b["image_url"]["url"].split(",")[1],
                    }
                }
                for b in blocks
            ]
            messages.append({"role": role, "parts": parts})

    if user_message.strip():
        if model.startswith("gpt"):
            messages.append({"role": "user", "content": user_message})
        else:
            messages.append(
                {"role": "user", "parts": [{"text": user_message}]})

    return {"messages": messages, "system_instruction": system_instruction}


async def build_gpt_message_blocks(
        user_message: str, files: List[UploadFile]) -> List[dict]:
    """Создаёт text/image-блоки для GPT-моделей на основе сообщения и файлов."""
    blocks: List[dict] = []

    if user_message := user_message.strip():
        blocks.append({"type": "text", "text": user_message})

    for f in files:
        name, ext = f.filename, Path(f.filename).suffix.lower()

        if ext in DOC_EXT:
            parsed = await parse_file(f)
            if parsed:
                blocks.append(
                    {"type": "text", "text": f"[File: {name}]\n{parsed}"})

        elif ext in IMG_EXT:
            b64 = base64.b64encode(await f.read()).decode()
            blocks.append({"type": "image_url", "image_url": {
                          "url": f"data:image/jpeg;base64,{b64}"}})

        else:
            blocks.append({"type": "text",
                           "text": f"[File: {name}] Unknown format, skipped."})

    return blocks


# ==============================
# БЛОК: CONTEXTENTRY – snapshot + KB-структура
# ==============================


IMAGE_HINT = (
    "You are given an image. Carefully describe everything that can be seen "
    "and convert the description into a short Q/A style block "
    "(3-6 concise questions and answers, no extra chit-chat)."
)


# async def ensure_entry_processed(
#         entry: "ContextEntry", ai_model: str = settings.MODEL_DEFAULT) -> None:
#     """
#     Гарантирует, что у ContextEntry заполнены snapshot_text и kb_structure.
#     При необходимости вызывает LLM.
#     """
#     updated: Dict[str, Any] = {}

#     # ---------- 1. Snapshot ----------
#     if entry.snapshot_text is None:
#         if entry.type is ContextType.URL and entry.url:
#             entry.snapshot_text = await cache_url_snapshot(str(entry.url))

#         elif entry.type is ContextType.FILE and entry.file_path:
#             ext = Path(entry.file_path).suffix.lower()
#             if ext in DOC_EXT:
#                 entry.snapshot_text = await parse_file_from_path(entry.file_path)
#             elif ext in IMG_EXT:
#                 entry.snapshot_text = f"[Image: {Path(entry.file_path).name}]"

#         elif entry.type is ContextType.TEXT and entry.text:
#             entry.snapshot_text = entry.text

#         if isinstance(entry.snapshot_text, bytes):
#             entry.snapshot_text = entry.snapshot_text.decode(errors="ignore")

#         if entry.snapshot_text:
#             updated["context.$.snapshot_text"] = entry.snapshot_text

#     # ---------- 2. KB-structure ----------
#     if entry.kb_structure is None:
#         blocks = await collect_context_blocks([entry])

#         # для картинок добавляем подсказку
#         if entry.type is ContextType.FILE and entry.file_path and Path(
#                 entry.file_path).suffix.lower() in IMG_EXT:
#             blocks.insert(0, {"sender_role": "user", "text": IMAGE_HINT})

#         if blocks:
#             sys_prompt = AI_PROMPTS["kb_structure_prompt"]
#             client, real_model = pick_model_and_client(ai_model)

#             msg = build_messages_for_model(sys_prompt, blocks, "", real_model)
#             raw = await admin_chat_generate_any(client, real_model, msg["messages"], msg.get("system_instruction"))

#             entry.kb_structure = extract_json_from_gpt(raw) or {}
#             updated["context.$.kb_structure"] = entry.kb_structure

#             # если KB-структура по изображению получена — делаем snapshot
#             # кратким QA
#             if entry.snapshot_text and entry.snapshot_text.startswith(
#                     "[Image:") and entry.kb_structure:
#                 flat: List[str] = []
#                 for t in entry.kb_structure.values():
#                     for s in t.get("subtopics", {}).values():
#                         for q, a in s.get("questions", {}).items():
#                             flat.append(f"Q: {q}\nA: {a.get('text', '')}\n")
#                 if flat:
#                     entry.snapshot_text = "\n".join(flat)
#                     updated["context.$.snapshot_text"] = entry.snapshot_text

#     # ---------- 3. Write-back ----------
#     if updated:
#         await mongo_db.knowledge_collection.update_one(
#             {"app_name": settings.APP_NAME, "context.id": entry.id},
#             {"$set": updated},
#         )

async def ensure_entry_processed(
    entry: "ContextEntry", ai_model: str = settings.MODEL_DEFAULT
) -> None:
    """
    Гарантирует, что у ContextEntry заполнены актуальные snapshot_text и kb_structure.
    При необходимости вызывает LLM.
    """
    updated: Dict[str, Any] = {}

    # ---------- 1. Snapshot ----------
    should_refresh_snapshot = False

    if entry.type == ContextType.URL and entry.url:
        # всегда обновляем snapshot (обновление кеша зависит от TTL)
        snapshot_text = await cache_url_snapshot(str(entry.url))
        if snapshot_text != entry.snapshot_text:
            entry.snapshot_text = snapshot_text
            updated["context.$.snapshot_text"] = entry.snapshot_text
            should_refresh_snapshot = True

    elif entry.snapshot_text is None:
        if entry.type == ContextType.FILE and entry.file_path:
            ext = Path(entry.file_path).suffix.lower()
            if ext in DOC_EXT:
                entry.snapshot_text = await parse_file_from_path(entry.file_path)
            elif ext in IMG_EXT:
                entry.snapshot_text = f"[Image: {Path(entry.file_path).name}]"

        elif entry.type == ContextType.TEXT and entry.text:
            entry.snapshot_text = entry.text

        if isinstance(entry.snapshot_text, bytes):
            entry.snapshot_text = entry.snapshot_text.decode(errors="ignore")

        if entry.snapshot_text:
            updated["context.$.snapshot_text"] = entry.snapshot_text
            should_refresh_snapshot = True

    # ---------- 2. KB-structure ----------
    if should_refresh_snapshot or entry.kb_structure is None:
        blocks = await collect_context_blocks([entry])

        if entry.type == ContextType.FILE and entry.file_path and Path(
            entry.file_path).suffix.lower() in IMG_EXT:
            blocks.insert(0, {"sender_role": "user", "text": IMAGE_HINT})

        if blocks:
            sys_prompt = AI_PROMPTS["kb_structure_prompt"]
            client, real_model = pick_model_and_client(ai_model)
            msg = build_messages_for_model(sys_prompt, blocks, "", real_model)

            raw = await admin_chat_generate_any(
                client, real_model, msg["messages"], msg.get("system_instruction")
            )

            new_kb = extract_json_from_gpt(raw) or {}
            if new_kb != entry.kb_structure:
                entry.kb_structure = new_kb
                updated["context.$.kb_structure"] = entry.kb_structure

                # обновляем краткий текст, если это было изображение
                if entry.snapshot_text and entry.snapshot_text.startswith("[Image:"):
                    flat: List[str] = []
                    for t in entry.kb_structure.values():
                        for s in t.get("subtopics", {}).values():
                            for q, a in s.get("questions", {}).items():
                                flat.append(f"Q: {q}\nA: {a.get('text', '')}\n")
                    if flat:
                        entry.snapshot_text = "\n".join(flat)
                        updated["context.$.snapshot_text"] = entry.snapshot_text

    # ---------- 3. Write-back ----------
    if updated:
        await mongo_db.knowledge_collection.update_one(
            {"app_name": settings.APP_NAME, "context.id": entry.id},
            {"$set": updated},
        )


# ---------- Вспомогательные функции для контекста ----------

async def collect_context_blocks(
    ctx_list: List["ContextEntry"],
    purposes: Optional[Set] = None,
) -> List[dict]:
    """Преобразует ContextEntry в text / image_url блоки для LLM."""
    allowed = purposes or {
        ContextPurpose.BOT,
        ContextPurpose.BOTH,
        ContextPurpose.KB,
        ContextPurpose.NONE,
    }

    blocks: List[dict] = []

    for e in ctx_list:
        if e.purpose not in allowed:
            continue

        if e.type == ContextType.TEXT and e.text:
            blocks.append({"sender_role": "Context", "text": e.text})

        elif e.type == ContextType.URL and e.url:
            txt = e.snapshot_text or await cache_url_snapshot(str(e.url))
            blocks.append({"sender_role": "Context", "text": txt})

        elif e.type == ContextType.FILE and e.file_path:
            ext = Path(e.file_path).suffix.lower()

            if ext in IMG_EXT:  # image -> image_url
                mime = mimetypes.guess_type(e.file_path)[0] or "image/jpeg"
                async with aiofiles.open(e.file_path, "rb") as f:
                    b64 = base64.b64encode(await f.read()).decode()
                blocks.append(
                    {"sender_role": "Context", "image_url": {
                        "url": f"data:{mime};base64,{b64}"}}
                )

            elif ext in DOC_EXT:  # doc/pdf/xls -> text
                txt = await parse_file_from_path(e.file_path)
                if txt:
                    blocks.append({"sender_role": "Context", "text": txt})

    return blocks


async def get_context_entry_content(entry: "ContextEntry") -> str:
    """Возвращает текст/снимок содержимого ContextEntry."""
    if entry.type is ContextType.TEXT:
        return entry.text or ""
    if entry.type is ContextType.FILE and entry.file_path:
        return await parse_file_from_path(entry.file_path)
    if entry.type is ContextType.URL and entry.url:
        return entry.snapshot_text or await cache_url_snapshot(str(entry.url))
    return ""

# ==============================
# БЛОК: HIGH-LEVEL –  анализ, патчи, KB-структуры
# ==============================


def normalize_snippets_structure(result: dict) -> List[dict]:
    """Нормализация структуры сниппетов."""
    norm = []
    for t in result.get("topics", []):
        if not isinstance(t, dict) or not isinstance(t.get("topic"), str):
            continue
        subs = []
        for s in t.get("subtopics", []):
            if isinstance(s, str):
                subs.append({"subtopic": s, "questions": {}})
            elif isinstance(s, dict) and isinstance(s.get("subtopic"), str):
                qs = s.get("questions", {})
                if isinstance(qs, list):
                    qs = {q: {"text": "", "files": []}
                          for q in qs if isinstance(q, str)}
                elif isinstance(qs, dict):
                    qs = {
                        q: {
                            "text": (a.get("text") if isinstance(a, dict) else ""),
                            "files": (a.get("files") if isinstance(a, dict) and isinstance(a.get("files"), list) else []),
                        }
                        for q, a in qs.items()
                        if isinstance(q, str)
                    }
                subs.append({"subtopic": s["subtopic"], "questions": qs})
        norm.append({"topic": t["topic"], "subtopics": subs})
    return norm


async def analyze_update_snippets_via_gpt(
    user_blocks: List[dict],
    user_info: str,
    knowledge_base: dict,
    bot_context_prompt: str,
    ai_model: str = settings.MODEL_DEFAULT,
) -> dict:
    """
    Определяет релевантные темы/подтемы, которые пользователь затронул,
    чтобы сократить KB-контекст для генерации патча.
    """
    _, kb_model = await get_knowledge_base()
    ctx_blocks = await collect_context_blocks(kb_model.context, {ContextPurpose.BOTH, ContextPurpose.KB})

    sys_prompt = AI_PROMPTS["analyze_update_snippets_prompt"].format(
        kb_full_structure=json.dumps(
            {
                t: {
                    "subtopics": {
                        s: {"questions": list(sub["questions"])}
                        for s, sub in topic["subtopics"].items()
                    }
                }
                for t, topic in knowledge_base.items()
            },
            ensure_ascii=False,
        ),
        bot_snippets_text=bot_context_prompt,
    )

    msg_bundle = build_messages_for_model(
        sys_prompt, ctx_blocks + user_blocks, "", ai_model)
    client, real_model = pick_model_and_client(ai_model)
    raw = await admin_chat_generate_any(client, real_model, msg_bundle["messages"], msg_bundle.get("system_instruction"))
    return {"topics": normalize_snippets_structure(extract_json_from_gpt(raw))}


async def extract_knowledge_with_structure(
    topics: List[dict],
    knowledge_base: dict,
) -> dict:
    """Извлекает из БЗ подмножество вопросов по заданной структуре."""
    extracted: dict = {}

    for itm in topics:
        t_name = itm.get("topic", "")
        if t_name not in knowledge_base:
            continue

        extracted.setdefault(t_name, {"subtopics": {}})
        topic_data = knowledge_base[t_name].get("subtopics", {})
        subs_required = itm.get("subtopics", [])

        if subs_required:
            for sub in subs_required:
                s_name = sub.get("subtopic")
                qs = sub.get("questions", [])
                if s_name and s_name in topic_data:
                    extracted[t_name]["subtopics"].setdefault(
                        s_name, {"questions": {}})
                    sub_data = topic_data[s_name]
                    if qs:
                        for q in qs:
                            if q in sub_data.get("questions", {}):
                                extracted[t_name]["subtopics"][s_name]["questions"][q] = sub_data["questions"][q]
                    else:
                        extracted[t_name]["subtopics"][s_name]["questions"] = sub_data.get(
                            "questions", {}).copy()
        else:
            for s_name, sub_d in topic_data.items():
                extracted[t_name]["subtopics"].setdefault(
                    s_name, {"questions": {}})
                extracted[t_name]["subtopics"][s_name]["questions"] = sub_d.get(
                    "questions", {}).copy()

    return extracted


async def generate_patch_body_via_gpt(
    user_blocks: List[dict],
    user_info: str,
    knowledge_base: dict,
    ai_model: str = settings.MODEL_DEFAULT,
) -> dict:
    """Генерирует JSON-патч для базы знаний."""
    _, kb_model = await get_knowledge_base()

    ctx_blocks = await collect_context_blocks(kb_model.context, {ContextPurpose.BOTH, ContextPurpose.KB})

    _, bot_prompt = await collect_bot_context_snippets(kb_model.context)

    snippets = await analyze_update_snippets_via_gpt(
        user_blocks=user_blocks,
        user_info=user_info,
        knowledge_base=knowledge_base,
        bot_context_prompt=bot_prompt,
        ai_model=ai_model,
    )

    kb_snippets_text = await extract_knowledge_with_structure(snippets["topics"], knowledge_base)

    sys_prompt = AI_PROMPTS["patch_body_prompt"].format(
        bot_snippets_text=bot_prompt,
        kb_snippets_text=kb_snippets_text,
    )

    msg_bundle = build_messages_for_model(
        sys_prompt, ctx_blocks + user_blocks, "", ai_model)
    client, real_model = pick_model_and_client(ai_model)
    raw = await admin_chat_generate_any(client, real_model, msg_bundle["messages"], msg_bundle.get("system_instruction"))

    return extract_json_from_gpt(raw)


async def build_kb_structure_from_context_entry(
    entry: ContextEntry, ai_model: str = settings.MODEL_DEFAULT
) -> dict:
    """KB-структура для одного ContextEntry (если ещё не сгенерирована)."""
    if getattr(entry, "kb_structure", None):
        return entry.kb_structure

    await ensure_entry_processed(entry, ai_model)
    return entry.kb_structure or {}


async def collect_bot_context_snippets(
        entries: List[ContextEntry]) -> Tuple[List[str], str]:
    """Текстовые сниппеты (purpose = KB/BOTH) + промпт."""
    frags = [
        (await get_context_entry_content(e)).strip()
        for e in entries
        if e.purpose in {ContextPurpose.KB, ContextPurpose.BOTH}
    ]
    frags = [f for f in frags if f]

    prompt = (
        "The following are text fragments that serve as background context for the assistant.\n"
        "They provide additional information that may be used to improve answers\n"
        "when relevant to the user's request, but should not be quoted directly or included\n"
        "verbatim in structured outputs.\n\n"
        + ("\n\n".join(frags) if frags else "No additional context available.")
    )

    return frags, prompt


async def collect_kb_structures_from_context(
    entries: List[ContextEntry],
    ai_model: str = settings.MODEL_DEFAULT,
) -> Tuple[List[dict], str]:
    """Собирает KB-структуры из списка контекста."""
    structs, blocks = [], []

    for e in entries:
        if e.purpose in {ContextPurpose.BOT, ContextPurpose.BOTH}:
            s = await build_kb_structure_from_context_entry(e, ai_model)
            if s:
                structs.append(s)
                blocks.append(
                    f"From: **{e.title}**\n"
                    f"```json\n{json.dumps(s, indent=2, ensure_ascii=False)}\n```"
                )

    prompt = (
        "The following structured knowledge blocks are extracted from source content\n"
        "and intended to be integrated into the main knowledge base.\n\n"
        + "\n\n".join(blocks)
    )
    return structs, prompt
